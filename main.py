import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import openpyxl
from docx import Document
from datetime import datetime
import os

def generate_from_template(template_path, data_map):
    doc = Document(template_path)

    for p in doc.paragraphs:
        for key, value in data_map.items():
            if key in p.text:
                p.text = p.text.replace(key, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in data_map.items():
                        if key in p.text:
                            p.text = p.text.replace(key, str(value))

    return doc


class ContainerTool:
    def __init__(self, root):
        self.root = root
        self.root.title("Tool Xuất Phiếu Container - Auto Export")
        self.root.geometry("800x600")

        # CHỐNG DUPLICATE
        self.is_running = False

        # Đường dẫn file
        self.ccx_file = ""
        self.daily_file = ""
        self.output_folder = ""

        # Cột cố định
        self.config = {
            'ccx_container_col': 4,
            'ccx_line_col': 3,
            'ccx_regid_col': 6,
            'daily_container_col': 1,
            'daily_iso_col': 6,
            'daily_damage_col': 7,
        }

        self.setup_ui()

    def setup_ui(self):
        title = tk.Label(self.root, text="XUẤT PHIẾU TỰ ĐỘNG",
                         font=("Arial", 18, "bold"), fg="#0066cc")
        title.pack(pady=20)

        subtitle = tk.Label(self.root,
                            text="Chọn 2 file Excel và tự động xuất tất cả phiếu",
                            font=("Arial", 10, "italic"))
        subtitle.pack()

        main_frame = ttk.Frame(self.root, padding="30")
        main_frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(main_frame, text="📄 BƯỚC 1: File Copy of CCX.xlsx",
                  font=("Arial", 12, "bold"),
                  foreground="#0066cc").grid(row=0, column=0, columnspan=3, sticky=tk.W, pady=(0, 10))

        ttk.Label(main_frame,
                  text="Chứa: Container, Line (Chủ KT), Reg ID",
                  font=("Arial", 9, "italic"),
                  foreground="gray").grid(row=1, column=0, columnspan=3, sticky=tk.W, padx=20)

        self.ccx_label = ttk.Label(main_frame, text="❌ Chưa chọn file",
                                   foreground="red", font=("Arial", 9))
        self.ccx_label.grid(row=2, column=0, columnspan=2, sticky=tk.W, padx=20, pady=5)

        ttk.Button(main_frame, text="📁 Chọn File",
                   command=self.select_ccx_file).grid(row=2, column=2, sticky=tk.E)

        ttk.Label(main_frame, text="📄 BƯỚC 2: File Daily Report.xlsx",
                  font=("Arial", 12, "bold"), foreground="#0066cc").grid(row=3, column=0, columnspan=3, sticky=tk.W, pady=(30, 10))

        ttk.Label(main_frame,
                  text="Chứa: Container, ISO, Damage Description",
                  font=("Arial", 9, "italic"),
                  foreground="gray").grid(row=4, column=0, columnspan=3, sticky=tk.W, padx=20)

        self.daily_label = ttk.Label(main_frame, text="❌ Chưa chọn file",
                                     foreground="red", font=("Arial", 9))
        self.daily_label.grid(row=5, column=0, columnspan=2, sticky=tk.W, padx=20, pady=5)

        ttk.Button(main_frame, text="📁 Chọn File",
                   command=self.select_daily_file).grid(row=5, column=2, sticky=tk.E)

        ttk.Label(main_frame, text="📁 BƯỚC 3: Thư mục lưu phiếu",
                  font=("Arial", 12, "bold"), foreground="#0066cc").grid(row=6, column=0, columnspan=3, sticky=tk.W, pady=(30, 10))

        self.output_label = ttk.Label(main_frame, text="❌ Chưa chọn thư mục",
                                      foreground="red", font=("Arial", 9))
        self.output_label.grid(row=7, column=0, columnspan=2, sticky=tk.W, padx=20, pady=5)

        ttk.Button(main_frame, text="📁 Chọn Thư Mục",
                   command=self.select_output_folder).grid(row=7, column=2, sticky=tk.E)

        export_btn = tk.Button(main_frame, text="🚀 XUẤT TẤT CẢ PHIẾU",
                               command=self.auto_export_all,
                               bg="#28a745", fg="white",
                               font=("Arial", 14, "bold"),
                               height=2, width=30)
        export_btn.grid(row=8, column=0, columnspan=3, pady=20)

        self.status_label = ttk.Label(main_frame,
                                      text="✓ Sẵn sàng - Hãy chọn 2 file Excel",
                                      relief=tk.SUNKEN,
                                      anchor=tk.W,
                                      font=("Arial", 9))
        self.status_label.grid(row=9, column=0, columnspan=3, sticky=(tk.W, tk.E))

    def select_ccx_file(self):
        file = filedialog.askopenfilename(
            title="Chọn file Copy of CCX",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        if file:
            self.ccx_file = file
            self.ccx_label.config(text=f"✓ {os.path.basename(file)}", foreground="green")

    def select_daily_file(self):
        file = filedialog.askopenfilename(
            title="Chọn file Daily Report",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        if file:
            self.daily_file = file
            self.daily_label.config(text=f"✓ {os.path.basename(file)}", foreground="green")

    def select_output_folder(self):
        folder = filedialog.askdirectory(title="Chọn thư mục lưu phiếu")
        if folder:
            self.output_folder = folder
            self.output_label.config(text=f"✓ {folder}", foreground="green")

    def read_ccx_containers(self):
        wb = openpyxl.load_workbook(self.ccx_file, data_only=True)

        ws = wb["STOCK"] if "STOCK" in wb.sheetnames else wb[wb.sheetnames[0]]

        containers = []

        for row in ws.iter_rows(min_row=4, values_only=True):
            if row and row[self.config['ccx_container_col']]:
                container = ''.join(c for c in str(row[self.config['ccx_container_col']]).strip() if c.isalnum()).upper()

                line = row[self.config['ccx_line_col']] or "ATL"
                regid = row[self.config['ccx_regid_col']] or "246789"

                containers.append({
                    "container": container,
                    "line": str(line),
                    "regid": str(regid)
                })

        wb.close()
        return containers

    def find_daily_info(self, container_no):
        wb = openpyxl.load_workbook(self.daily_file, data_only=True)

        ws = wb["STOCK"] if "STOCK" in wb.sheetnames else wb[wb.sheetnames[0]]

        for row in ws.iter_rows(min_row=6, values_only=True):
            if row and row[self.config['daily_container_col']]:
                cell_container = ''.join(c for c in str(row[self.config['daily_container_col']]).strip() if c.isalnum()).upper()

                if cell_container == container_no:
                    iso = row[self.config['daily_iso_col']] or "2270"
                    damage = row[self.config['daily_damage_col']] or ""
                    wb.close()
                    return {"iso": str(iso), "damage": str(damage)}

        wb.close()
        return None

    def create_word_document(self, container_data):
        if hasattr(sys, "_MEIPASS"):
            template_path = os.path.join(sys._MEIPASS, "PHIEU_TEMPLATE.docx")
        else:
            template_path = "PHIEU_TEMPLATE.docx"

        today = datetime.now()

        data_map = {
            "{{REGID}}": container_data['regid'],
            "{{DATE}}": "",
            "{{TIME}}": "",
            "{{CONTAINER}}": container_data['container'],
            "{{ISO}}": container_data['iso'],
            "{{LINE}}": container_data['line'],
            "{{STATUS}}": "R",
            "{{DAMAGE}}": container_data['damage'],
        }

        return generate_from_template(template_path, data_map)

    def auto_export_all(self):

        # ============================
        #   KHÓA CHỐNG GỌI NHIỀU LẦN
        # ============================
        if self.is_running:
            print("DEBUG: auto_export_all bị gọi lại → CHẶN")
            return
        self.is_running = True
        # ============================

        if not self.ccx_file:
            messagebox.showerror("Lỗi", "❌ Vui lòng chọn file CCX!")
            self.is_running = False
            return

        if not self.daily_file:
            messagebox.showerror("Lỗi", "❌ Vui lòng chọn file Daily Report!")
            self.is_running = False
            return

        if not self.output_folder:
            messagebox.showerror("Lỗi", "❌ Vui lòng chọn thư mục lưu phiếu!")
            self.is_running = False
            return

        try:
            self.status_label.config(text="⏳ Đang đọc file CCX...")
            self.root.update()

            containers = self.read_ccx_containers()

            if not containers:
                messagebox.showwarning("Cảnh báo", "❌ Không tìm thấy container nào trong file CCX!")
                self.is_running = False
                return

            success = 0
            fail_list = []

            for idx, ccx_data in enumerate(containers):

                self.status_label.config(text=f"⏳ Đang xử lý {idx+1}/{len(containers)}: {ccx_data['container']}")
                self.root.update()

                daily_info = self.find_daily_info(ccx_data["container"])

                if not daily_info:
                    fail_list.append(f"❌ {ccx_data['container']} - Không tìm thấy trong Daily")
                    continue

                full_data = {
                    "container": ccx_data["container"],
                    "line": ccx_data["line"],
                    "regid": ccx_data["regid"],
                    "iso": daily_info["iso"],
                    "damage": daily_info["damage"]
                }

                doc = self.create_word_document(full_data)

                filename = f"PHIEU_{ccx_data['container']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                filepath = os.path.join(self.output_folder, filename)
                doc.save(filepath)

                success += 1

            msg = f"🎉 HOÀN THÀNH!\n\n"
            msg += f"✅ Thành công: {success}/{len(containers)} phiếu\n"

            if fail_list:
                msg += f"\n❌ Thất bại: {len(fail_list)} container\n"
                for item in fail_list[:10]:
                    msg += f"   {item}\n"
                if len(fail_list) > 10:
                    msg += f"... và {len(fail_list)-10} lỗi khác\n"

            msg += f"\n📁 File lưu tại:\n{self.output_folder}"

            messagebox.showinfo("Kết quả", msg)
            self.status_label.config(text=f"✓ Hoàn thành! Xuất {success} phiếu")

        except Exception as e:
            messagebox.showerror("Lỗi", f"❌ Có lỗi xảy ra:\n{str(e)}")

        finally:
            self.is_running = False


def main():
    root = tk.Tk()
    app = ContainerTool(root)
    root.mainloop()


if __name__ == "__main__":
    main()
